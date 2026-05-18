"""Claude Sonnet resume tailor — superior reasoning for strategic resume rewrites."""
from __future__ import annotations

import asyncio
import logging
import re
from pathlib import Path

from docx import Document
from temporalio import activity

from shared.config import settings
from shared.llm_client import claude_chat
from shared.models import JobPost, MatchResult, UserProfile

log = logging.getLogger("worker.resume")

RESUME_DIR = Path("/app/profiles")

TAILOR_PROMPT = """\
You are a world-class resume writer specialising in technical roles.
Rewrite the candidate's resume to precisely match this job description.
Your goal is to maximise interview callback rate by strategically framing
the candidate's real experience in the exact language, priorities, and keywords
the hiring team is looking for. Surface the most relevant work prominently.
Do NOT invent experience — reframe and emphasise what is already there.

JOB:
Company: {company}
Title: {title}
Description:
{description}

CANDIDATE SKILLS: {skills}

BASE RESUME (numbered paragraph format):
{resume_content}

RULES — follow every one:
- Rewrite every [NORMAL] and [BULLET] paragraph to mirror the JD's exact language, \
action verbs, and keywords
- Open the summary with the job title or closest equivalent, then immediately lead \
with the most relevant skills in the JD's vocabulary
- For each bullet: identify the JD's top 3 requirements, then make sure at least one \
bullet per role directly addresses each requirement using the JD's own phrasing
- Reorder bullets ONLY within the same [ROLE] section — never move content across roles
- Leave [HEADER], [ROLE], and [EMPTY] lines structurally unchanged (light wording ok for [ROLE])
- Never invent facts — every claim must be traceable to the original
- Preserve all numbers, dollar amounts, dates, and company names exactly
- Return the EXACT same number of lines as the input
- Output ONLY the numbered paragraph list, no commentary

OUTPUT FORMAT (same as input):
1|[HEADER] Name
2|[NORMAL] Contact info
...
"""


def _read_docx_structured(docx_path: Path) -> str:
    """Extract numbered, tagged paragraph list from a DOCX file."""
    doc = Document(str(docx_path))
    lines = []
    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            lines.append(f"{i}|[EMPTY]")
            continue
        style = para.style.name.lower()
        is_heading = "heading" in style
        is_bold = any(r.bold for r in para.runs if r.text.strip())
        is_allcaps = text.isupper() and len(text) > 3
        is_bullet = "list" in style or text.startswith(("•", "-", "–", "*"))
        # Detect role lines: title|company|dates pattern
        is_role = bool(re.search(r"\d{4}", text) and "|" in text)

        if is_role:
            tag = "[ROLE]"
        elif is_heading or is_bold or is_allcaps:
            tag = "[HEADER]"
        elif is_bullet:
            tag = "[BULLET]"
        else:
            tag = "[NORMAL]"

        lines.append(f"{i}|{tag} {text}")
    return "\n".join(lines)


def _apply_rewrites_to_docx(original_path: Path, output_path: Path, rewritten: str) -> None:
    """Apply rewritten paragraph text back onto the original DOCX template."""
    rewrites: dict[int, str] = {}
    for line in rewritten.splitlines():
        line = line.strip()
        if not line or "|" not in line:
            continue
        num_str, rest = line.split("|", 1)
        try:
            num = int(num_str.strip())
        except ValueError:
            continue
        text = re.sub(r"^\[(HEADER|NORMAL|BULLET|ROLE|EMPTY)\]\s*", "", rest).strip()
        if "[EMPTY]" not in rest:
            rewrites[num] = text

    import shutil
    shutil.copy2(str(original_path), str(output_path))
    doc = Document(str(output_path))

    for idx, para in enumerate(doc.paragraphs, start=1):
        if idx not in rewrites:
            continue
        new_text = rewrites[idx]
        if not new_text:
            continue
        # Preserve formatting of first run, clear the rest
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.add_run(new_text)

    doc.save(str(output_path))


def _download_from_minio(minio_path: str, person_id: str, variant_id: str) -> Path:
    """Download a resume template from MinIO to /tmp. Raises if not found."""
    from minio import Minio
    from minio.error import S3Error

    client = Minio(
        settings.minio_endpoint,
        access_key=settings.minio_access_key,
        secret_key=settings.minio_secret_key,
        secure=settings.minio_secure,
    )
    dest = Path("/tmp") / "resumes" / person_id / f"{variant_id}.docx"
    dest.parent.mkdir(parents=True, exist_ok=True)
    try:
        client.fget_object(settings.minio_bucket, minio_path, str(dest))
        log.info("Downloaded resume from MinIO: %s → %s", minio_path, dest)
        return dest
    except S3Error as e:
        raise FileNotFoundError(
            f"Resume template not in MinIO at '{minio_path}'. "
            f"Run: python scripts/upload_resumes.py from the Mac to seed MinIO. "
            f"Original error: {e}"
        ) from e


@activity.defn
async def tailor_resume(job_dict: dict, profile_dict: dict, match_dict: dict) -> str:
    """
    Tailor resume for a specific job. Returns the local path of the generated DOCX.
    The storage activity will upload it to MinIO.
    """
    job = JobPost(**job_dict)
    profile = UserProfile(**profile_dict)
    match = MatchResult(**match_dict)

    # Find the right resume variant.
    # match.recommended_variant may be a bare key like "infra" or "aiml"; try an exact
    # match first, then a suffix match (e.g. "infra" matches "sai_infra").
    raw_hint = match.recommended_variant or ""
    variant = next(
        (v for v in profile.resume_variants if v.id == raw_hint), None
    ) or next(
        (v for v in profile.resume_variants if v.id.endswith(f"_{raw_hint}") or v.id == raw_hint),
        profile.resume_variants[0] if profile.resume_variants else None,
    )

    if not variant:
        raise RuntimeError(f"No resume variant found for {profile.id}")

    variant_id = variant.id  # always use the canonical ID from the profile

    # Resume is stored locally at /app/profiles/{person_id}/resumes/{variant_id}.docx
    docx_path = RESUME_DIR / profile.id / "resumes" / f"{variant_id}.docx"
    if not docx_path.exists():
        # Fall back: try downloading the base template from MinIO
        minio_path = variant.minio_path
        log.info("Local resume not found — downloading from MinIO: %s", minio_path)
        docx_path = _download_from_minio(minio_path, profile.id, variant_id)

    structured = _read_docx_structured(docx_path)

    prompt = TAILOR_PROMPT.format(
        company=job.company,
        title=job.title,
        description=job.description_text[:4000],
        skills=", ".join(profile.skills),
        resume_content=structured,
    )
    rewritten = await asyncio.to_thread(
        claude_chat,
        prompt,
        model=settings.claude_model,
        max_tokens=4000,
    )

    # Build output path
    safe_company = re.sub(r"[^\w-]", "-", job.company)
    safe_title = re.sub(r"[^\w-]", "-", job.title)
    output_filename = f"{profile.id}_{safe_company}_{safe_title}_resume.docx"
    output_path = Path("/tmp") / output_filename

    _apply_rewrites_to_docx(docx_path, output_path, rewritten)
    log.info("Resume tailored for %s @ %s → %s", job.title, job.company, output_path)
    return str(output_path)
