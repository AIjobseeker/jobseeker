#!/usr/bin/env python3
"""Demo: tailor resume + generate cover letter for ONE job, locally,
without Temporal / postgres / minio. Uses Ollama directly.

Saves all artifacts under ~/.jobseeker/docs/<dedup_id>/ and (optionally)
sends them as Telegram attachments threaded to a single match alert.

  python3 scripts/demo_tailor.py                           # generate only
  python3 scripts/demo_tailor.py --send-telegram           # generate + send
  python3 scripts/demo_tailor.py --output /tmp/jobseeker_demo

What it produces (under <output> AND ~/.jobseeker/docs/<dedup_id>/):
  job.json                  the job dict that was tailored against
  resume.tailored.docx      DOCX with bullets rewritten to mirror the JD
  cover_letter.txt          3-4 paragraph cover letter
  study_guide.md            interview prep guide

It also persists `resume_path`, `cover_letter_path`, `study_guide` into
seen.db keyed by dedup_id, so the mobile app / next runs can find them.

Use this for: showing someone what the tool produces. Not the production
pipeline (that runs continuously via Temporal worker).
"""
from __future__ import annotations

import argparse
import asyncio
import hashlib
import json
import os
import re
import shutil
import sqlite3
import sys
from datetime import datetime
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO))

# Load .env so OLLAMA_HOST etc are available
try:
    from dotenv import load_dotenv

    load_dotenv(REPO / ".env", override=False)
except ImportError:
    pass


def _short_dedup_id(company: str, source_id: str) -> str:
    raw = f"{company.lower()}|{source_id}".encode("utf-8")
    return hashlib.sha256(raw).hexdigest()[:16]


# ── Sample job (used when --job-index is not provided and no seen.db) ─────
SAMPLE_JOB = {
    "id": "demo-uuid",
    "source_id": "demo-001",
    "source": "greenhouse",
    "company": "Stripe",
    "title": "Staff Site Reliability Engineer, Platform",
    "description_text": (
        "Stripe is hiring a Staff SRE on the Platform Reliability team. "
        "You'll lead our Kubernetes platform serving production workloads "
        "across multiple regions, drive SLO/error-budget discipline, and "
        "make decisions about infrastructure-as-code (Terraform), CI/CD, "
        "and observability (Prometheus, Grafana, distributed tracing). "
        "8+ years of experience operating complex distributed systems in "
        "production. Strong fluency in Python or Go for tooling. "
        "Experience leading incident response. We sponsor H1B transfers."
    ),
    "url": "https://stripe.com/jobs/staff-sre-platform",
    "location": "San Francisco, CA / Remote (US)",
    "department": "Infrastructure",
    "remote": True,
    "scraped_at": "2026-05-17T18:00:00Z",
}


def _green(s):  return f"\033[32m{s}\033[0m"
def _yellow(s): return f"\033[33m{s}\033[0m"
def _bold(s):   return f"\033[1m{s}\033[0m"
def _dim(s):    return f"\033[2m{s}\033[0m"


def pick_job(seen_db: Path, override_index: int | None) -> dict:
    """Prefer the highest-scoring real job from seen.db; fall back to sample."""
    if override_index is not None:
        return SAMPLE_JOB
    if not seen_db.exists():
        return SAMPLE_JOB
    try:
        conn = sqlite3.connect(str(seen_db))
        row = conn.execute(
            "SELECT key, company, title, url, score FROM seen_jobs "
            "WHERE notified=1 ORDER BY score DESC LIMIT 1"
        ).fetchone()
        conn.close()
        if not row:
            return SAMPLE_JOB
        # We only have summary fields; fabricate a description that's good
        # enough for tailoring. In production the description text would
        # come from the original scrape (it's in the NDJSON publish but not
        # stored in seen.db).
        _key, company, title, url, score = row
        return {
            **SAMPLE_JOB,
            "company": company, "title": title, "url": url,
            "description_text": (
                f"Real job from your last scrape (score={score:.2f}). "
                f"Ollama will use only this title + your profile to tailor. "
                f"For richer tailoring, we should pipe the full JD text from "
                f"the scraper through to the dedup DB."
            ),
        }
    except Exception:
        return SAMPLE_JOB


def load_profile() -> dict:
    parsed = REPO / "profiles" / "sai" / "profile.parsed.yaml"
    if not parsed.exists():
        print(_yellow(f"WARN: {parsed} not found — using a stub profile."))
        return {
            "person": {"name": "Saikrishna Narvaneni"},
            "seniority": {"level": "staff"},
            "core_skills": ["kubernetes", "terraform", "aws", "python", "linux", "ci/cd"],
            "target_titles": ["sre", "platform engineer", "staff sre"],
            "profile_summary": "Staff SRE / Platform engineer with 8 years of experience.",
        }
    import yaml
    return yaml.safe_load(parsed.read_text())


# ── Resume parsing/rewriting ─────────────────────────────────────────────

def docx_to_structured(docx_path: Path) -> tuple[str, dict[int, str]]:
    """Return (numbered_text_for_LLM, original_paragraph_map).

    Paragraphs get [HEADER]/[NORMAL]/[BULLET]/[ROLE]/[EMPTY] tags. The LLM
    is instructed to rewrite bullets/normals while leaving structure alone.
    """
    from docx import Document

    doc = Document(str(docx_path))
    lines: list[str] = []
    originals: dict[int, str] = {}
    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        originals[i] = text
        if not text:
            lines.append(f"{i}|[EMPTY]")
            continue
        style = para.style.name.lower() if para.style and para.style.name else ""
        is_heading = "heading" in style
        is_bullet = "list" in style or text.startswith(("•", "-", "–", "*"))
        is_role = bool(re.search(r"\b\d{4}\b", text)) and ("|" in text or " - " in text or "—" in text)
        if is_role:
            tag = "[ROLE]"
        elif is_heading or (text.isupper() and len(text) > 3):
            tag = "[HEADER]"
        elif is_bullet:
            tag = "[BULLET]"
        else:
            tag = "[NORMAL]"
        lines.append(f"{i}|{tag} {text}")
    return "\n".join(lines), originals


def apply_rewrites(src_docx: Path, out_docx: Path, rewritten: str) -> int:
    """Write the LLM's tailored paragraphs back into a copy of src_docx.
    Returns the count of paragraphs changed.
    """
    from docx import Document

    rewrites: dict[int, str] = {}
    for line in rewritten.splitlines():
        line = line.strip()
        if not line or "|" not in line:
            continue
        num_str, _, rest = line.partition("|")
        try:
            num = int(num_str.strip())
        except ValueError:
            continue
        # strip leading [TAG] from the rewritten text
        body = re.sub(r"^\[(HEADER|NORMAL|BULLET|ROLE|EMPTY)\]\s*", "", rest).strip()
        if "[EMPTY]" in rest or not body:
            continue
        rewrites[num] = body

    shutil.copy2(str(src_docx), str(out_docx))
    doc = Document(str(out_docx))
    changed = 0
    for idx, para in enumerate(doc.paragraphs, start=1):
        if idx not in rewrites:
            continue
        new_text = rewrites[idx]
        # Preserve formatting of first run, clear the rest
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.add_run(new_text)
        changed += 1
    doc.save(str(out_docx))
    return changed


# ── Ollama prompts ───────────────────────────────────────────────────────

TAILOR_PROMPT = """You are a world-class technical-resume writer.

Rewrite the candidate's resume to align with the job below — without inventing facts.

JOB:
Company: {company}
Title: {title}
JD:
{jd}

CANDIDATE'S PROFILE:
{profile_summary}
Core skills: {core_skills}

BASE RESUME (numbered paragraphs with tags):
{resume}

RULES:
- Rewrite [BULLET] and [NORMAL] paragraphs to mirror the JD's vocabulary and action verbs.
- Reorder bullets ONLY within their own [ROLE] section.
- Leave [HEADER], [ROLE], and [EMPTY] lines structurally unchanged.
- NEVER invent skills, employers, dates, or numbers. Every claim must trace to the original.
- Preserve all numbers, dollar amounts, dates, and company names exactly.
- Output the EXACT same line numbering as the input.
- Output ONLY the numbered paragraph list. No commentary.
"""

COVER_PROMPT = """You are an expert career coach writing a tight, specific cover letter.

JOB:
Company: {company}
Title: {title}
JD:
{jd}

CANDIDATE:
Name: {name}
Years: {years}
Top skills aligned to this JD: {skills}

RULES:
- 3-4 paragraphs, under 350 words.
- Open with a specific hook about the company or role. NEVER start with "I am excited to apply".
- Paragraph 2: 2 concrete achievements that map to the JD's needs (no metrics fabrication).
- Paragraph 3: why THIS company specifically.
- Confident IC tone. No filler ("passionate", "team player", "results-driven", "dream company").
- Output ONLY the letter body. No subject, date, or address headers.
"""

STUDY_PROMPT = """You are a senior technical interviewer.

Job: {title} at {company}
JD key points: {jd}
Candidate strong areas: {strong}

Output a markdown interview-prep guide. Max 300 words. Sections:

## Must Know
- 3-5 technical topics central to this JD

## Brush Up
- 2-3 things present but worth refreshing

## Likely Interview Questions
- 3 system-design or technical questions

## Quick Tips
- 1-2 company-specific items to research
"""


async def ollama_chat(model: str, prompt: str, max_tokens: int = 2000, think: bool = False) -> str:
    from ollama import AsyncClient

    host = os.environ.get("OLLAMA_HOST", "http://100.115.111.9:11434")
    client = AsyncClient(host=host)
    resp = await client.chat(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        think=think,
        options={"num_predict": max_tokens},
    )
    return (resp.message.content or "").strip()


async def claude_chat(prompt: str, max_tokens: int = 4000) -> str:
    """Use Claude via Anthropic API or an internal corporate proxy.

    Looks at ANTHROPIC_BASE_URL / ANTHROPIC_AUTH_TOKEN / ANTHROPIC_API_KEY
    in env. Same flexibility as scripts/parse_resume.py.
    """
    import anthropic

    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip() or "dummy"
    base_url = os.environ.get("ANTHROPIC_BASE_URL", "").strip() or None
    auth_token = os.environ.get("ANTHROPIC_AUTH_TOKEN", "").strip()

    extra_headers: dict[str, str] = {}
    if auth_token:
        extra_headers["Authorization"] = f"Bearer {auth_token}"
    raw_extra = os.environ.get("ANTHROPIC_EXTRA_HEADERS_JSON", "").strip()
    if raw_extra:
        try:
            extra_headers.update(json.loads(raw_extra))
        except json.JSONDecodeError:
            pass

    kwargs: dict = {"api_key": api_key}
    if base_url:
        kwargs["base_url"] = base_url
    if extra_headers:
        kwargs["default_headers"] = extra_headers
    client = anthropic.Anthropic(**kwargs)
    model = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6")
    resp = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(b.text for b in resp.content if getattr(b, "type", "") == "text")


async def llm_chat(provider: str, model: str, prompt: str, max_tokens: int, think: bool) -> str:
    """Dispatch to ollama or claude based on provider."""
    if provider == "claude":
        return await claude_chat(prompt, max_tokens=max_tokens)
    return await ollama_chat(model, prompt, max_tokens=max_tokens, think=think)


# ── Telegram delivery ────────────────────────────────────────────────────

async def telegram_send_message(token: str, chat_id: str, text: str,
                                reply_markup: dict | None = None) -> int | None:
    """Send a markdown message; returns the new message_id or None."""
    import httpx

    body = {
        "chat_id": chat_id,
        "text": text,
        "parse_mode": "Markdown",
        "disable_web_page_preview": True,
    }
    if reply_markup is not None:
        body["reply_markup"] = reply_markup
    async with httpx.AsyncClient(timeout=20) as c:
        r = await c.post(
            f"https://api.telegram.org/bot{token}/sendMessage", json=body,
        )
    if r.status_code != 200:
        print(f"  WARN: telegram sendMessage HTTP {r.status_code}: {r.text[:200]}")
        return None
    try:
        return int(r.json()["result"]["message_id"])
    except (KeyError, ValueError, TypeError):
        return None


async def telegram_send_document(token: str, chat_id: str, file_path: Path,
                                 caption: str, reply_to: int | None) -> bool:
    """Upload a file as a Telegram document. Returns True on success."""
    import httpx

    with open(file_path, "rb") as f:
        async with httpx.AsyncClient(timeout=60) as c:
            data = {
                "chat_id": chat_id,
                "caption": caption,
                "parse_mode": "Markdown",
            }
            if reply_to is not None:
                data["reply_to_message_id"] = str(reply_to)
            files = {"document": (file_path.name, f, "application/octet-stream")}
            r = await c.post(
                f"https://api.telegram.org/bot{token}/sendDocument",
                data=data, files=files,
            )
    if r.status_code != 200:
        print(f"  WARN: telegram sendDocument HTTP {r.status_code}: {r.text[:200]}")
        return False
    return True


def _build_alert_text(job: dict, profile: dict, study_preview: str) -> str:
    """The single 'NEW MATCH' alert message that wraps everything together."""
    company = job["company"]
    title = job["title"]
    location = job.get("location", "Location TBD")
    score = job.get("_score")
    score_line = f"Score: {int(round(score * 100))}% | " if score is not None else ""
    name = profile.get("person", {}).get("name", "you")
    return (
        f"*NEW MATCH — {company}*\n"
        f"*{title}*\n"
        f"{score_line}{location}\n"
        f"\n"
        f"Tailored resume + cover letter attached below.\n"
        f"\n"
        f"*Interview prep:*\n"
        f"```\n{study_preview[:600]}\n```\n"
        f"\n"
        f"_Generated for {name} at {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC_"
    )


def _alert_keyboard(dedup_id: str, apply_url: str) -> dict:
    return {
        "inline_keyboard": [
            [
                {"text": "Apply (open job)", "url": apply_url},
                {"text": "Mark Applied", "callback_data": f"applied:{dedup_id}"},
            ],
            [
                {"text": "Skip", "callback_data": f"skip:{dedup_id}"},
                {"text": "Save for Later", "callback_data": f"saved:{dedup_id}"},
            ],
        ]
    }


# ── Main ─────────────────────────────────────────────────────────────────

async def main() -> int:
    ap = argparse.ArgumentParser(description="Demo resume tailor + cover letter")
    ap.add_argument("--output", type=Path, default=Path("/tmp/jobseeker_demo"))
    ap.add_argument("--job-index", type=int, default=None,
                    help="Force the sample job (otherwise pulls top from seen.db)")
    ap.add_argument("--resume", type=Path,
                    default=REPO / "profiles" / "sai" / "resumes" / "sai_infra.docx")
    ap.add_argument("--match-model", default="qwen3:14b")
    ap.add_argument("--cover-model", default="qwen3:8b")
    ap.add_argument("--study-model", default=os.environ.get("OLLAMA_MODEL_STUDY", "qwen2.5:7b"))
    ap.add_argument("--use-claude", action="store_true",
                    help="Use Claude (Anthropic / internal proxy) for resume + cover letter "
                         "instead of Ollama. Better output quality. Reads the same env "
                         "vars parse_resume.py uses.")
    ap.add_argument("--send-telegram", action="store_true",
                    help="After generating, send to Telegram with attachments")
    ap.add_argument("--no-persist", action="store_true",
                    help="Skip writing to seen.db (default: persist)")
    args = ap.parse_args()

    args.output.mkdir(parents=True, exist_ok=True)

    print(_bold(f"\n=== JobSeeker tailor demo ==="))
    print(f"  output dir:   {args.output}")
    print(f"  base resume:  {args.resume}")
    print(f"  ollama host:  {os.environ.get('OLLAMA_HOST', 'http://100.115.111.9:11434')}")

    if not args.resume.exists():
        print(_yellow(f"\nERROR: base resume not found at {args.resume}"))
        print(_yellow(f"Place a DOCX there or pass --resume. Exiting."))
        return 1

    # Pick a job from seen.db (matches your live pipeline) or the sample.
    seen_db_default = Path("~/.jobseeker/seen.db").expanduser()
    job = pick_job(seen_db_default, args.job_index)
    dedup_id = _short_dedup_id(job["company"], job["source_id"])
    print(f"  dedup_id:     {_yellow(dedup_id)}")

    profile = load_profile()
    name = profile.get("person", {}).get("name", "Saikrishna Narvaneni")
    summary = profile.get("profile_summary", "Staff SRE / Platform engineer.")
    core_skills = profile.get("core_skills", []) or []
    if isinstance(core_skills, list) and core_skills and isinstance(core_skills[0], dict):
        core_skills = [s.get("name", "") for s in core_skills]
    years = profile.get("seniority", {}).get("min_years") or 8

    (args.output / "job.json").write_text(json.dumps(job, indent=2))
    print(f"\n{_bold('Job picked:')}  {_green(job['company'])}  {job['title']}")
    print(f"  -> {args.output / 'job.json'}")

    structured, originals = docx_to_structured(args.resume)
    n_paragraphs = len(originals)
    print(f"\n{_bold('Step 1')}  parsing base DOCX")
    print(f"  {n_paragraphs} paragraphs read from {args.resume.name}")

    # ── Tailor resume ─────────────────────────────────────────────────────
    tailor_provider = "claude" if args.use_claude else "ollama"
    tailor_label = (
        f"{_yellow('Claude (Sonnet 4.6)')}" if args.use_claude
        else f"{_yellow(args.match_model)} (think=ON)"
    )
    print(f"\n{_bold('Step 2')}  tailoring resume via {tailor_label}")
    tailored = await llm_chat(
        provider=tailor_provider,
        model=args.match_model,
        prompt=TAILOR_PROMPT.format(
            company=job["company"],
            title=job["title"],
            jd=job["description_text"][:4000],
            profile_summary=summary,
            core_skills=", ".join(core_skills[:15]),
            resume=structured,
        ),
        max_tokens=4000,
        think=True,
    )
    out_docx = args.output / "resume.tailored.docx"
    changed = apply_rewrites(args.resume, out_docx, tailored)
    print(f"  {_green(str(changed))} of {n_paragraphs} paragraphs rewritten")
    print(f"  -> {out_docx}")
    (args.output / "resume.tailored.txt").write_text(tailored)

    # ── Cover letter ─────────────────────────────────────────────────────
    cover_label = (
        f"{_yellow('Claude (Sonnet 4.6)')}" if args.use_claude
        else f"{_yellow(args.cover_model)} (think=OFF, fast)"
    )
    print(f"\n{_bold('Step 3')}  cover letter via {cover_label}")
    cover = await llm_chat(
        provider=tailor_provider,
        model=args.cover_model,
        prompt=COVER_PROMPT.format(
            company=job["company"],
            title=job["title"],
            jd=job["description_text"][:3000],
            name=name,
            years=years,
            skills=", ".join(core_skills[:8]),
        ),
        max_tokens=900,
        think=False,
    )
    cover_path = args.output / "cover_letter.txt"
    cover_path.write_text(cover)
    print(f"  -> {cover_path}")

    # ── Study guide ───────────────────────────────────────────────────────
    print(f"\n{_bold('Step 4')}  interview prep via {_yellow(args.study_model)}")
    study = await ollama_chat(
        args.study_model,
        STUDY_PROMPT.format(
            title=job["title"],
            company=job["company"],
            jd=job["description_text"][:2000],
            strong=", ".join(core_skills[:10]),
        ),
        max_tokens=600,
    )
    study_path = args.output / "study_guide.md"
    study_path.write_text(study)
    print(f"  -> {study_path}")

    # ── Persist artifacts (DB + canonical store) ─────────────────────────
    canonical = None
    if not args.no_persist:
        from services.notifier.artifact_store import get_default_store
        from services.notifier.dedup import DedupStore, compute_key, TAILOR_DONE

        store = get_default_store()
        # Canonical filenames make it obvious what's what when files are
        # listed flat (e.g. in MinIO). Filename ends up in Telegram caption.
        safe_company = re.sub(r"[^\w-]", "-", job["company"])[:40]
        safe_title = re.sub(r"[^\w-]", "-", job["title"])[:40]
        resume_name = f"resume_{safe_company}_{safe_title}.docx"
        cover_name = f"cover_letter_{safe_company}_{safe_title}.txt"

        resume_canonical = store.put_file(dedup_id, out_docx, name=resume_name)
        cover_canonical = store.put_file(dedup_id, cover_path, name=cover_name)
        # study guide: small enough to inline in DB but also keep on disk
        store.put_text(dedup_id, "study_guide.md", study)

        canonical = {
            "resume": (store.resolve_local(resume_canonical), resume_canonical),
            "cover":  (store.resolve_local(cover_canonical), cover_canonical),
        }

        seen_db = Path(os.environ.get("JOBSEEKER_DB_PATH", "~/.jobseeker/seen.db")).expanduser()
        with DedupStore(seen_db) as db:
            full_key = compute_key(job["company"], job["source_id"])
            # Insert the row first if it doesn't exist (so demo works on a
            # fresh db without requiring the live pipeline to have seen this job).
            db.insert_if_new(
                key=full_key, company=job["company"], title=job["title"],
                url=job["url"], score=job.get("_score") or 0.0,
            )
            db.record_artifacts(
                key=full_key,
                resume_path=resume_canonical,
                cover_letter_path=cover_canonical,
                study_guide=study,
                tailor_status=TAILOR_DONE,
            )
        print(f"\n{_bold('Step 5')}  artifacts persisted")
        print(f"  store dir:    {store.dir_for(dedup_id)}")
        print(f"  seen.db:      {seen_db} (dedup_id={dedup_id})")
        print(f"  resume_path:  {resume_canonical}")
        print(f"  cover_path:   {cover_canonical}")

    # ── Optional: send to Telegram ───────────────────────────────────────
    if args.send_telegram:
        token = os.environ.get("TELEGRAM_BOT_TOKEN", "").strip()
        chat = (os.environ.get("TELEGRAM_CHAT_ID_SAI")
                or os.environ.get("TELEGRAM_CHAT_ID", "")).strip()
        if not token or not chat:
            print(_yellow(
                "\nWARN: TELEGRAM_BOT_TOKEN or TELEGRAM_CHAT_ID_SAI not set "
                "(skipping Telegram send). Source .env in the same shell."
            ))
        else:
            print(f"\n{_bold('Step 6')}  sending to Telegram (chat={chat})")
            alert_text = _build_alert_text(job, profile, study)
            keyboard = _alert_keyboard(dedup_id, job["url"])
            msg_id = await telegram_send_message(token, chat, alert_text, keyboard)
            if msg_id is None:
                print(_yellow("  alert send failed — nothing to thread to"))
            else:
                print(f"  alert sent: msg_id={msg_id}")
                # Threaded replies — appear directly under the alert in chat.
                ok1 = await telegram_send_document(
                    token, chat, out_docx,
                    caption=f"Tailored resume — *{job['company']}*", reply_to=msg_id,
                )
                ok2 = await telegram_send_document(
                    token, chat, cover_path,
                    caption=f"Cover letter — *{job['company']}*", reply_to=msg_id,
                )
                print(f"  resume sent:       {_green('yes') if ok1 else _yellow('no')}")
                print(f"  cover letter sent: {_green('yes') if ok2 else _yellow('no')}")

                # Persist the message_id so bot_listener can edit it later.
                if not args.no_persist:
                    from services.notifier.dedup import DedupStore, compute_key
                    seen_db = Path(os.environ.get(
                        "JOBSEEKER_DB_PATH", "~/.jobseeker/seen.db"
                    )).expanduser()
                    with DedupStore(seen_db) as db:
                        full_key = compute_key(job["company"], job["source_id"])
                        db.mark_notified(full_key, message_id=msg_id)

    # ── Console preview ─────────────────────────────────────────────────
    print(f"\n{_bold('=' * 70)}")
    print(_bold("  Cover letter preview"))
    print(_bold("=" * 70))
    print(_dim(cover[:1200]))
    if len(cover) > 1200:
        print(_dim(f"...[{len(cover) - 1200} more chars in {cover_path}]"))

    print(f"\n{_bold('=' * 70)}")
    print(_bold("  Study guide preview"))
    print(_bold("=" * 70))
    print(_dim(study[:1200]))

    print(f"\n{_bold('=' * 70)}")
    print(f"All files in: {_green(str(args.output))}")
    print(f"  open {args.output}/resume.tailored.docx       # see DOCX in Word/Pages")
    print(f"  cat  {args.output}/cover_letter.txt")
    print(f"  cat  {args.output}/study_guide.md")
    if not args.no_persist:
        store_dir = Path("~/.jobseeker/docs").expanduser() / dedup_id
        print(f"\nCanonical store: {store_dir}")
        print(f"DB query:        sqlite3 ~/.jobseeker/seen.db "
              f"\"SELECT resume_path, cover_letter_path FROM seen_jobs "
              f"WHERE substr(key,1,16)='{dedup_id}'\"")
    print(_bold("=" * 70))
    return 0


if __name__ == "__main__":
    sys.exit(asyncio.run(main()))
