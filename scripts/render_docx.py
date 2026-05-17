"""DOCX renderer for tailored resumes.

The same plan + resume that builds the HTML drives the DOCX. Why a separate
module? `tailor_v2.py` is already 900+ lines and mixing python-docx layout
with the HTML template makes both harder to follow.

DOCX is what most ATS systems prefer for upload (HTML can be rejected,
PDFs sometimes get parsed as a single text blob). We give the user all
three: HTML for browser preview, PDF for sharing, DOCX for ATS upload.
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional


def render_docx(
    *,
    resume: dict,
    plan: dict,
    out_path: Path,
    show_highlights: bool = True,
) -> None:
    """Build a clean, ATS-friendly single-column DOCX.

    `plan` is the same JSON the LLM emits (rewrites or selects). `resume` is
    the source resume.yaml. We mirror the HTML template's section ordering:
      Header -> Summary -> Highlights (optional) -> Experience -> Skills -> Certs -> Edu.

    No "Tailored for ..." footer — that's a recruiter-visible AI tell. The
    same metadata lives in match_report.json next to this file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Tight, ATS-friendly margins — wider than 0.5" causes Workday to choke
    # on parsing some line-spacing combos. 0.7" is the sweet spot.
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ── Default font (Helvetica/Arial — ATS-safe, not Calibri-only) ──
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(10)

    person = resume.get("person", {}) or {}

    # Name
    p = doc.add_paragraph()
    run = p.add_run(person.get("name", ""))
    run.bold = True
    run.font.size = Pt(18)
    p.paragraph_format.space_after = Pt(2)

    # Contact line
    contact_parts = []
    for k in ("email", "phone", "location", "linkedin", "github"):
        v = person.get(k)
        if v:
            contact_parts.append(str(v))
    p = doc.add_paragraph()
    run = p.add_run(" · ".join(contact_parts))
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_after = Pt(6)

    # ── Summary ──
    summary_text = (plan.get("summary") or "").strip()
    if summary_text:
        _add_section_heading(doc, "Summary")
        p = doc.add_paragraph(summary_text)
        p.paragraph_format.space_after = Pt(4)

    # ── Highlights (optional) ──
    highlights = resume.get("key_highlights") or []
    if show_highlights and highlights:
        _add_section_heading(doc, "Selected Highlights")
        for h in highlights[:6]:
            _add_bullet(doc, h)

    # ── Experience ──
    _add_section_heading(doc, "Experience")
    picks_by_idx = {p.get("role_index"): p for p in (plan.get("experience") or [])}
    for ridx, role in enumerate(resume.get("experience", []) or [], start=1):
        sel = picks_by_idx.get(ridx)
        bullets_to_show: list[str] = []
        bullets = role.get("bullets", []) or []
        if sel and sel.get("rewrites"):
            for r in sel["rewrites"]:
                t = (r.get("rewritten") or "").strip()
                if t:
                    bullets_to_show.append(t)
        elif sel:
            keep = [i for i in (sel.get("keep_bullet_indices") or []) if 1 <= i <= len(bullets)]
            bullets_to_show = [bullets[i - 1] for i in keep]
        else:
            bullets_to_show = bullets[:3]

        if not bullets_to_show:
            continue

        # role line: "Company · Title — Location | Date"
        p = doc.add_paragraph()
        head = p.add_run(f"{role.get('company', '')} · {role.get('title', '')}")
        head.bold = True
        head.font.size = Pt(10.5)
        date_range = f"{role.get('start', '')} - {role.get('end', '')}".strip(" -")
        loc = role.get("location") or ""
        sub_parts = [x for x in [loc, date_range] if x]
        if sub_parts:
            sub = doc.add_paragraph(" | ".join(sub_parts))
            sub.paragraph_format.space_after = Pt(2)
            sr = sub.runs[0]
            sr.italic = True
            sr.font.size = Pt(9.5)
            sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for b in bullets_to_show:
            _add_bullet(doc, b)

    # ── Skills ──
    skills = resume.get("skills")
    if skills:
        _add_section_heading(doc, "Skills")
        if isinstance(skills, dict):
            for group, items in skills.items():
                if not items:
                    continue
                p = doc.add_paragraph()
                gn = p.add_run(f"{group}: ")
                gn.bold = True
                p.add_run(", ".join(str(s) for s in items))
                p.paragraph_format.space_after = Pt(1)
        else:
            doc.add_paragraph(", ".join(str(s) for s in skills))

    # ── Certifications ──
    certs = resume.get("certifications") or []
    if certs:
        _add_section_heading(doc, "Certifications")
        for c in certs:
            if isinstance(c, str):
                doc.add_paragraph(c)
            else:
                line = c.get("name", "")
                if c.get("year"):
                    line = f"{line} ({c['year']})"
                if c.get("issuer"):
                    line = f"{line} — {c['issuer']}"
                doc.add_paragraph(line)

    # ── Education ──
    edu = resume.get("education") or []
    if edu:
        _add_section_heading(doc, "Education")
        for e in edu:
            line = e.get("degree", "")
            if e.get("field"):
                line = f"{line}, {e['field']}"
            line = f"{line} — {e.get('school', '')}"
            if e.get("end"):
                line = f"{line} ({e['end']})"
            doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _add_section_heading(doc, text: str) -> None:
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def _add_bullet(doc, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    if p.runs:
        p.runs[0].text = text
        p.runs[0].font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
