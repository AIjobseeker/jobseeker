"""Generate GF's DOCX resume files from resume.yaml, then upload ALL resumes to MinIO.

Run once from Mac (where MinIO is accessible and DOCX files live):
  cd /path/to/jobseeker
  python scripts/gen_and_upload_resumes.py

After this the Windows worker can pull resume templates from MinIO automatically.
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from minio import Minio
from minio.error import S3Error

PROFILES_DIR = Path(__file__).parent.parent / "profiles"

# MinIO config — read from env or use defaults
MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "localhost:9000")
MINIO_ACCESS = os.getenv("MINIO_ACCESS_KEY", "minioadmin")
MINIO_SECRET = os.getenv("MINIO_SECRET_KEY", "minioadmin123")
MINIO_BUCKET = os.getenv("MINIO_BUCKET", "jobseeker-docs")


# ── DOCX generation ──────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_bold_line(doc: Document, text: str, font_size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)


def _add_normal(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _build_gf_docx(resume: dict, archetype: str) -> Document:
    """Build a styled DOCX for a given archetype variant."""
    p = resume["person"]
    doc = Document()

    # ── Header ──
    _add_heading(doc, p["name"], level=1)
    _add_normal(doc, f"{p['email']}  |  {p['phone']}  |  {p['location']}")
    _add_normal(doc, p["linkedin"])
    doc.add_paragraph()

    # ── Summary ──
    _add_heading(doc, "SUMMARY", level=2)
    _add_normal(doc, resume["summary"].strip())
    doc.add_paragraph()

    # ── Key Highlights (archetype-specific ordering) ──
    _add_heading(doc, "KEY HIGHLIGHTS", level=2)
    for hl in resume["key_highlights"]:
        _add_bullet(doc, hl)
    doc.add_paragraph()

    # ── Experience ──
    _add_heading(doc, "EXPERIENCE", level=2)
    for exp in resume["experience"]:
        role_line = f"{exp['title']}  |  {exp['company']}  |  {exp['location']}  |  {exp['start']} – {exp['end']}"
        _add_bold_line(doc, role_line)
        for b in exp["bullets"]:
            _add_bullet(doc, b)
        doc.add_paragraph()

    # ── Projects ──
    _add_heading(doc, "AI / ML PROJECTS", level=2)
    for proj in resume.get("ai_ml_projects", []):
        proj_title = proj["name"]
        if "publication" in proj:
            proj_title += f"  ({proj['publication']})"
        _add_bold_line(doc, proj_title)
        for b in proj["bullets"]:
            _add_bullet(doc, b)
        doc.add_paragraph()

    # ── Skills (group by section, show archetype-relevant first) ──
    _add_heading(doc, "SKILLS", level=2)
    skills = resume.get("skills", {})

    # Priority order depends on archetype
    if archetype == "aiml":
        group_order = ["ML / AI Frameworks", "ML Domains", "LLM Engineering",
                       "Audio & Signal Processing", "ML Practice & Evaluation",
                       "Languages", "Data & Visualization", "MLOps & Serving",
                       "Web & Frontend", "Cloud & DevOps", "Databases"]
    elif archetype == "data":
        group_order = ["Data & Visualization", "ML Practice & Evaluation",
                       "ML / AI Frameworks", "Languages", "Databases",
                       "MLOps & Serving", "LLM Engineering", "Cloud & DevOps"]
    else:  # swe
        group_order = ["Languages", "Web & Frontend", "Cloud & DevOps", "Databases",
                       "ML / AI Frameworks", "MLOps & Serving", "Data & Visualization"]

    shown = set()
    for group in group_order:
        if group in skills:
            _add_bold_line(doc, group + ":", font_size=10)
            _add_normal(doc, "  " + ", ".join(skills[group]))
            shown.add(group)
    # Any remaining groups
    for group, items in skills.items():
        if group not in shown:
            _add_bold_line(doc, group + ":", font_size=10)
            _add_normal(doc, "  " + ", ".join(items))
    doc.add_paragraph()

    # ── Education ──
    _add_heading(doc, "EDUCATION", level=2)
    for edu in resume.get("education", []):
        _add_bold_line(doc, f"{edu['degree']} in {edu['field']}")
        _add_normal(doc, f"{edu['school']}, {edu['location']}  |  {edu['start']} – {edu['end']}")
    doc.add_paragraph()

    # ── Certifications ──
    certs = resume.get("certifications", [])
    if certs:
        _add_heading(doc, "CERTIFICATIONS", level=2)
        for cert in certs:
            issuer = f" — {cert['issuer']}" if cert.get("issuer") else ""
            _add_bullet(doc, f"{cert['name']}{issuer}")
        doc.add_paragraph()

    # ── Honors ──
    honors = resume.get("honors", [])
    if honors:
        _add_heading(doc, "HONORS & LEADERSHIP", level=2)
        for h in honors:
            _add_bullet(doc, h)
        doc.add_paragraph()

    # ── Publications ──
    pubs = resume.get("publications", [])
    if pubs:
        _add_heading(doc, "PUBLICATIONS", level=2)
        for pub in pubs:
            _add_normal(doc, f"{pub['title']} — {pub['venue']} ({pub['year']})")
            _add_normal(doc, pub.get("url", ""))
        doc.add_paragraph()

    return doc


def generate_gf_docs() -> list[Path]:
    resume_yaml = PROFILES_DIR / "gf" / "resume.yaml"
    with open(resume_yaml) as f:
        resume = yaml.safe_load(f)

    resumes_dir = PROFILES_DIR / "gf" / "resumes"
    resumes_dir.mkdir(parents=True, exist_ok=True)

    variants = [
        ("gf_aiml", "aiml"),
        ("gf_swe", "swe"),
        ("gf_data", "data"),
    ]
    generated = []
    for variant_id, archetype in variants:
        out = resumes_dir / f"{variant_id}.docx"
        doc = _build_gf_docx(resume, archetype)
        doc.save(str(out))
        print(f"  Generated: {out}")
        generated.append(out)
    return generated


# ── MinIO upload ─────────────────────────────────────────────────────────────

def upload_all_to_minio() -> None:
    client = Minio(MINIO_ENDPOINT, access_key=MINIO_ACCESS, secret_key=MINIO_SECRET, secure=False)

    if not client.bucket_exists(MINIO_BUCKET):
        client.make_bucket(MINIO_BUCKET)
        print(f"  Created bucket: {MINIO_BUCKET}")

    docx_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # Collect all DOCX files in profiles/*/resumes/
    for docx_path in sorted(PROFILES_DIR.glob("*/resumes/*.docx")):
        person_id = docx_path.parent.parent.name
        variant_id = docx_path.stem
        minio_key = f"resumes/{person_id}/{variant_id}.docx"
        client.fput_object(MINIO_BUCKET, minio_key, str(docx_path), content_type=docx_type)
        print(f"  Uploaded: {docx_path.name}  →  minio://{MINIO_BUCKET}/{minio_key}")


def main():
    print("=== Generating GF resume DOCX files ===")
    generate_gf_docs()

    print("\n=== Uploading all resumes to MinIO ===")
    try:
        upload_all_to_minio()
        print("\nDone! Windows worker will auto-download templates from MinIO.")
    except Exception as e:
        print(f"\nMinIO upload failed: {e}")
        print("Make sure MinIO is running (docker compose up minio) and accessible.")
        sys.exit(1)


if __name__ == "__main__":
    main()
